import pytest
from pathlib import Path
from docx import Document
from src.converter import convert_markdown_to_word, correct_image_paths, process_markdown_files


# Create a temporary input directory for testing
@pytest.fixture
def input_dir(tmp_path):
    """Create a temporary image directory for testing."""
    image_dir = tmp_path / "images"
    image_dir.mkdir()
    return image_dir


# Create a temporary output directory for testing
@pytest.fixture
def output_dir(tmp_path):
    """Create a temporary output directory for testing."""
    output_dir = tmp_path / "output"
    output_dir.mkdir()
    return output_dir


# Test basic Markdown to Word conversion with proper content and structure
def test_basic_markdown_to_word_conversion(input_dir, output_dir):
    md_content = "# Heading 1\n\n## Heading 2\n\njust a [link](https://www.google.com)"
    md_file = input_dir / "test.md"
    md_file.write_text(md_content)
    
    output_file = output_dir / "test.docx"

    convert_markdown_to_word(md_file, output_file)

    assert output_file.exists(), "Output Word file was not created"
    assert output_file.stat().st_size > 0, "Output Word file is empty"

    doc = Document(output_file)
    assert len(doc.paragraphs) == 3, "Document should have 3 paragraphs"
    assert doc.paragraphs[0].text == "Heading 1", "First paragraph should be Heading 1"
    assert doc.paragraphs[1].text == "Heading 2", "Second paragraph should be Heading 2"
    assert "link" in doc.paragraphs[2].text, "Link text should be in the document"
    assert len(doc.inline_shapes) == 0, "Document should have no images"


# Test image path correction when local images are not available
def test_image_path_correction_for_missing_images(input_dir):
    md_content = "![alt text](https://example.com/missing.png)"
    corrected_content = correct_image_paths(md_content, input_dir)
    expected_content = "[CLICK TO VIEW ONLINE IMAGE: (alt text)](https://example.com/missing.png)"
    assert corrected_content == expected_content, "External links should be wrapped with 'CLICK TO VIEW ONLINE IMAGE' text"


# Test image path correction when local images are available
def test_image_path_correction_for_local_images(input_dir):
    md_file = input_dir / "test.md"
    md_file.touch()
    (input_dir / "local.jpg").touch()
    md_content = "![alt text](https://example.com/local.jpg)"
    corrected_content = correct_image_paths(md_content, md_file)
    assert corrected_content == f"![alt text]({input_dir / 'local.jpg'})", "Links should have local target"


# Test image path correction with a mix of local and missing images
def test_image_path_correction_for_mixed_local_and_missing(input_dir):
    (input_dir / "local.gif").touch()
    (input_dir / "local.png").touch()

    # Create a markdown file
    md_file = input_dir / "test.md"
    md_file.touch()

    md_content = """# Mixed Image Test
    ![](http://example.com/local.gif)
    ![](http://example.com/local.png)
    ![](http://example.com/nonexistent.gif)
    this is a [text link](http://somewhere.com/)
    """

    corrected_content = correct_image_paths(md_content, md_file)

    expected = f"""# Mixed Image Test
    ![]({input_dir / "local.gif"})
    ![]({input_dir / "local.png"})
    [CLICK TO VIEW ONLINE IMAGE: ()](http://example.com/nonexistent.gif)
    this is a [text link](http://somewhere.com/)
    """

    assert corrected_content == expected, "Incorrect handling of mixed local and missing images"


# Test processing of multiple Markdown files into Word documents
def test_multiple_markdown_files_processing(input_dir, output_dir):
    for i in range(3):
        md_file = input_dir / f"test{i}.md"
        md_file.write_text(f"# Test {i}\n![Image {i}](https://example.com/local{i}.png)")
        (input_dir / f"local{i}.png").touch()

    process_markdown_files(input_dir, output_dir)

    for i in range(3):
        output_file = output_dir / f"test{i}.docx"
        assert output_file.exists(), f"Word file test{i}.docx was not created"
        assert output_file.stat().st_size > 0, f"Word file test{i}.docx is empty"


# Test correct handling of local and online images in processed files
#   Verifies that local images are correctly included in the generated Word document,
#   while online images are referenced as text.
def test_local_and_online_images_in_processed_files(input_dir, output_dir):
    local_image = input_dir / "local_image.png"
    local_image.touch()
    md_content = f"""# Test
![Local Image]({local_image})
![Online Image](https://example.com/online_image.png)
"""
    md_file = input_dir / "test_mixed_images.md"
    md_file.write_text(md_content)

    process_markdown_files(input_dir, output_dir)
    output_file = output_dir / "test_mixed_images.docx"

    assert output_file.exists(), "Word file with mixed images was not created"
    assert output_file.stat().st_size > 0, "Word file with mixed images is empty"
    
    # Check Word document content
    doc = Document(output_file)
    assert len(doc.paragraphs) == 2, "Document should have two paragraphs"
    assert doc.paragraphs[0].text == "Test", "First paragraph should be the title"
    assert len(doc.inline_shapes) == 1, "Document should have one image (local)"
    assert "CLICK TO VIEW ONLINE IMAGE: (Online Image)" in doc.paragraphs[1].text, "Online image URL should be in the document"

    # NEEDED: An assert to ensure 
    # "CLICK TO VIEW ONLINE IMAGE: (Online Image)" is a hyperlink with target "https://example.com/online_image.png"
